import os

from PIL import Image
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm


class WordClient:
    def __init__(self):
        self.doc_name = 'Отчёт'
        self.doc = Document()

    def add_heading2(self, title: str, heading_no: int) -> None:
        self.doc.add_heading(
            f"{heading_no}. Решения задач на тему «{title}»",
            level=2,
        )
        self.doc.add_paragraph()

    def add_solution(self, no: int, title: str, descr: str, img_path: str) -> None:
        self.doc.add_paragraph(f"«{title}».")
        self.doc.add_paragraph(f"{descr}")
        self.doc.add_paragraph()

        if os.path.exists(img_path):
            img = Image.open(img_path)
            width_px, height_px = img.size

            max_width_mm = 140
            max_width_px = int(max_width_mm * 96 / 25.4)

            if width_px > max_width_px:
                ratio = max_width_px / width_px
                new_height = int(height_px * ratio)
                img = img.resize((max_width_px, new_height))
                img.save(img_path)

            pic_p = self.doc.add_paragraph()
            run = pic_p.add_run()
            run.add_picture(img_path, width=Mm(max_width_mm))
            pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            label = self.doc.add_paragraph(f"Рисунок {no} — решение задачи «{title}».")
            label.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            self.doc.add_paragraph("[Изображение не найдено]")

        self.doc.add_paragraph()

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, doc_name: str) -> None:
        os.makedirs("my_docs", exist_ok=True)
        if not doc_name.endswith(".docx"):
            doc_name += ".docx"
        self.doc.save(f"my_docs/{doc_name}")
